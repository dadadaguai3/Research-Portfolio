import os
import re
import json
import time
import pandas as pd
from datetime import datetime
from openai import OpenAI
import docx
import shutil

try:
    from win32com.client import Dispatch
except ImportError:
    Dispatch = None

# ==============================================================================
# 配置（已脱敏，请根据你的实际项目路径修改）
# ==============================================================================
ROOT_DIR = r"C:\Project\Data\Reports_Source"
EXCEL_CONFIG_PATH = r"C:\Project\Config\city_list.xlsx"
OUTPUT_EXCEL_PATH = r"C:\Project\Output\analysis_summary.xlsx"
BREAKPOINT_FILE = r"C:\Project\Config\breakpoint_record.json"
PATH_TEXT_FILE = r"C:\Project\Logs\city_path_list.txt"

# 实际上传文本的保存文件夹
PROMPT_LOG_DIR = r"C:\Project\Logs\AI_Prompt_Logs"

# 核心表另存文件夹
PRIMARY_FILES_SAVE_DIR = r"C:\Project\Output\Primary_Tables_Backup"

# 已经替换为占位符
QWEN_API_KEY = "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"
QWEN_MODEL = "qwen-plus"

SUPPORTED_EXTS = ['.xlsx', '.xls', '.docx', '.doc', '.txt', '.csv']

PARAMETERS = [
    "财政事务", "税收事务", "审计事务",
    "财政事务——信息化建设", "税收事务——信息化建设", "审计事务——信息化建设",
]

# 严格列规则
VALID_DEC_COLS = ["决算数", "决算金额", "决算支出"]
INVALID_BUD_COLS = ["预算数", "年初预算", "调整预算"]

# 优先读取的核心表关键词（增强版）
PRIMARY_TABLE_KEYWORDS = [
    "一般公共预算本级支出表",  # 最高优先级
    "一般公共预算支出决算表",
    "市本级一般公共预算支出",
    "一般公共预算支出决算",
    "市本级支出决算表",
    "一般公共预算支出"
]

# 仅保留市本级/市级关键词，排除全市
CITY_LEVEL_KEYWORDS = ["市级", "市本级", "本级"]
# 需要排除的关键词
EXCLUDE_KEYWORDS = ["全市"]

# 优先级关键词定义
SHI_BENJI_EXP_KEYWORD = "市本级支出决算表"  # 市本级支出决算表优先
SHI_JI_EXP_KEYWORD = "市级支出决算表"  # 市级支出决算表优先
SHI_BENJI_KEYWORDS = ["市本级"]  # 市本级通用关键词
SHI_JI_KEYWORDS = ["市级"]  # 市级通用关键词

# 工作表优先级关键词
SHEET_PRIORITY_KEYWORDS = {
    "市本级": ["市本级", "本级"],  # 最高优先级
    "市级": ["市级"],  # 第二优先级
    "排除": ["全市", "全地区", "汇总"]  # 需忽略的工作表
}


# ==============================================================================
# 主类
# ==============================================================================
class RobustBudgetAnalyzer:
    def __init__(self):
        self.client = OpenAI(api_key=QWEN_API_KEY, base_url="https://dashscope.aliyuncs.com/compatible-mode/v1")
        self.target_keywords = []
        self.breakpoints = self._load_breakpoints()
        self.city_config_df = self._load_full_city_config()
        self.city_order_list = self._get_city_order()

        # 创建核心表备份文件夹
        os.makedirs(PRIMARY_FILES_SAVE_DIR, exist_ok=True)
        # 创建AI交互文本日志文件夹
        os.makedirs(PROMPT_LOG_DIR, exist_ok=True)

    def _load_breakpoints(self):
        if os.path.exists(BREAKPOINT_FILE):
            try:
                with open(BREAKPOINT_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {"completed": []}
        return {"completed": []}

    def _save_breakpoint(self, city_name):
        if city_name not in self.breakpoints["completed"]:
            self.breakpoints["completed"].append(city_name)
            with open(BREAKPOINT_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.breakpoints, f, ensure_ascii=False, indent=2)

    def _load_full_city_config(self):
        if not os.path.exists(EXCEL_CONFIG_PATH):
            print(f"配置文件不存在：{EXCEL_CONFIG_PATH}")
            return pd.DataFrame()
        try:
            df = pd.read_excel(EXCEL_CONFIG_PATH, dtype=str).fillna("")
            if not all(c in df.columns for c in ["省名", "市名"]):
                return pd.DataFrame()
            if "是否跳过" not in df.columns:
                df["是否跳过"] = "否"
            df["市本级路径"] = df.apply(lambda r: os.path.join(ROOT_DIR, r["省名"], r["市名"], "市本级"), axis=1)
            df = df.drop_duplicates(subset=["市名"])
            self._save_paths_to_text(df)
            return df
        except Exception as e:
            print(f"加载配置文件出错：{e}")
            return pd.DataFrame()

    def _save_paths_to_text(self, df):
        try:
            with open(PATH_TEXT_FILE, 'w', encoding='utf-8') as f:
                f.write(f"生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}\n")
                f.write("-" * 80 + "\n")
                for _, r in df.iterrows():
                    f.write(f"{r['省名']} {r['市名']} | {r['市本级路径']}\n")
        except Exception as e:
            print(f"保存路径文件出错：{e}")

    def _get_city_order(self):
        if self.city_config_df.empty:
            return []
        return self.city_config_df[self.city_config_df["市名"] != ""]["市名"].tolist()

    def _is_city_skipped(self, city):
        if self.city_config_df.empty:
            return False
        row = self.city_config_df[self.city_config_df["市名"] == city]
        return row.iloc[0]["是否跳过"] == "是" if not row.empty else False

    def _get_city_path(self, city):
        if self.city_config_df.empty:
            return ""
        row = self.city_config_df[self.city_config_df["市名"] == city]
        return row.iloc[0]["市本级路径"] if not row.empty else ""

    # ==========================================================================
    # 判断是否 多工作表文件
    # ==========================================================================
    def _is_multi_sheet_file(self, file_path):
        """判断是否为 多工作表 Excel 文件"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in ['.xlsx', '.xls']:
            return False
        try:
            return len(pd.ExcelFile(file_path).sheet_names) >= 2
        except:
            return False

    # ==========================================================================
    # 新增：识别多表同工文件（单工作表包含多个表格的Excel文件）
    # ==========================================================================
    def _is_multi_table_single_sheet(self, file_path):
        """
        判断是否为多表同工文件：
        1. 仅针对Excel文件（xlsx/xls）
        2. 只有一个工作表
        3. 工作表内容包含多个独立表格（通过空行/空列分隔，或行数超过阈值）
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in ['.xlsx', '.xls']:
            return False

        try:
            # 1. 检查工作表数量
            excel_file = pd.ExcelFile(file_path)
            if len(excel_file.sheet_names) != 1:
                return False

            # 2. 读取唯一工作表的内容
            sheet_name = excel_file.sheet_names[0]
            df = pd.read_excel(file_path, sheet_name=sheet_name, dtype=str, header=None)
            df = df.fillna("")

            # 3. 判断是否包含多个表格的特征
            total_rows = len(df)
            empty_row_threshold = 5  # 连续空行阈值
            table_separators = 0
            empty_row_count = 0

            # 统计连续空行数（作为表格分隔符）
            for idx, row in df.iterrows():
                if all(cell == "" for cell in row.values):
                    empty_row_count += 1
                    if empty_row_count >= empty_row_threshold:
                        table_separators += 1
                        empty_row_count = 0
                else:
                    empty_row_count = 0

            # 判定条件：
            # - 存在至少1个表格分隔符 或
            # - 总行数超过200行（大概率包含多个表格）
            is_multi_table = table_separators >= 1 or total_rows > 200

            if is_multi_table:
                file_name = os.path.basename(file_path)
                if "决算" in file_name:
                    print(f"     识别为高优先级决算类多表同工文件: {file_name}")
                elif "预算" in file_name:
                    print(f"     识别为预算类多表同工文件: {file_name}")
            return is_multi_table

        except Exception as e:
            # 出错时默认不判定为多表同工文件
            return False

    # ==========================================================================
    # 统一判断：是否是【多表同工 或 多工作表】文件
    # ==========================================================================
    def _is_special_important_file(self, file_path):
        return self._is_multi_table_single_sheet(file_path) or self._is_multi_sheet_file(file_path)

    # ==========================================================================
    # 新增：判断多表同工文件类型（决算/预算）
    # ==========================================================================
    def _get_multi_table_type(self, file_path):
        """判断多表同工文件类型：决算/预算"""
        if not self._is_multi_table_single_sheet(file_path):
            return ""

        file_name = os.path.basename(file_path).lower()
        if "决算" in file_name:
            return "决算"
        elif "预算" in file_name:
            return "预算"
        return ""

    # ==========================================================================
    # 保存核心表到备份文件夹
    # ==========================================================================
    def _save_primary_file(self, source_path, city, year):
        """将识别的核心表另存到备份文件夹"""
        try:
            # 构造保存路径：城市_年份_原文件名
            file_name = os.path.basename(source_path)
            safe_city = re.sub(r'[\\/:*?"<>|]', '', city)
            save_name = f"{safe_city}_{year}_{file_name}"
            save_path = os.path.join(PRIMARY_FILES_SAVE_DIR, save_name)

            # 避免重名
            counter = 1
            while os.path.exists(save_path):
                name, ext = os.path.splitext(save_name)
                save_path = os.path.join(PRIMARY_FILES_SAVE_DIR, f"{name}_{counter}{ext}")
                counter += 1

            # 复制文件
            shutil.copy2(source_path, save_path)
            print(f"     核心表已备份: {save_name}")
            return save_path
        except Exception as e:
            print(f"     核心表备份失败: {e}")
            return

    # ==========================================================================
    # 检查文件是否包含 市本级/市级/本级（增强版：考虑城市名）
    # ==========================================================================
    def _has_city_keyword(self, file_path, city_name=""):
        """
        检查文件是否包含有效的市级关键词
        规则：包含城市名+（市本级/市级/本级）才有效
        只有城市名没有级别关键词的忽略
        """
        fname = os.path.basename(file_path).lower()

        # 提取城市名（去掉"市"字）
        city_core = city_name.replace("市", "") if city_name else ""

        # 检查文件名
        # 有效模式：城市名+市本级/市级/本级
        valid_patterns = [
            f"{city_core}市本级",
            f"{city_core}市级",
            f"{city_core}本级",
            f"{city_name}市本级",
            f"{city_name}市级",
            f"{city_name}本级"
        ]

        # 如果文件名直接包含有效模式，返回True
        for pattern in valid_patterns:
            if pattern.lower() in fname:
                return True

        # 如果只有城市名但没有级别关键词，返回False
        if city_core and city_core in fname:
            # 检查是否同时有级别关键词
            has_level = any(kw in fname for kw in ["市本级", "本级", "市级"])
            if not has_level:
                return False

        # 原有的检查（兜底）
        if any(kw in fname for kw in ["市本级", "本级", "市级"]):
            return True

        # 检查文件内容
        return self._check_file_content_for_city_level(file_path, city_name)

    # ==========================================================================
    # 检查文件内容是否包含市级/市本级关键词（增强版：考虑城市名）
    # ==========================================================================
    def _check_file_content_for_city_level(self, file_path, city_name=""):
        """检查文件内容的前三行是否包含城市名+市级/市本级关键词"""
        ext = os.path.splitext(file_path)[1].lower()
        city_core = city_name.replace("市", "") if city_name else ""

        # 构建需要匹配的模式
        target_patterns = [
            f"{city_core}市本级",
            f"{city_core}市级",
            f"{city_core}本级",
            f"{city_name}市本级",
            f"{city_name}市级",
            f"{city_name}本级"
        ]

        try:
            if ext in ['.xlsx', '.xls']:
                # 读取Excel的前几行
                df = pd.read_excel(file_path, sheet_name=None, nrows=3, dtype=str)
                for sheet_name, sheet_df in df.items():
                    # 将前3行转换为字符串检查
                    for _, row in sheet_df.head(3).iterrows():
                        row_text = ' '.join([str(val) for val in row.values if pd.notna(val)])
                        # 检查是否包含有效模式
                        for pattern in target_patterns:
                            if pattern in row_text:
                                return True
                        # 如果只有城市名没有级别关键词，返回False
                        if city_core and city_core in row_text:
                            has_level = any(kw in row_text for kw in CITY_LEVEL_KEYWORDS)
                            if not has_level:
                                continue

            elif ext in ['.docx', '.doc']:
                # 读取Word文档的前几段
                tmp = file_path
                is_tmp = False
                if ext == '.doc' and Dispatch:
                    tmp = self._doc2docx(file_path)
                    is_tmp = True
                if not tmp:
                    return False

                doc = docx.Document(tmp)
                # 检查前10个段落
                for i, para in enumerate(doc.paragraphs[:10]):
                    for pattern in target_patterns:
                        if pattern in para.text:
                            if is_tmp and os.path.exists(tmp):
                                os.remove(tmp)
                            return True

                # 检查第一个表格的前几行
                if doc.tables:
                    first_table = doc.tables[0]
                    for i, row in enumerate(first_table.rows[:3]):
                        row_text = ' '.join([cell.text for cell in row.cells])
                        for pattern in target_patterns:
                            if pattern in row_text:
                                if is_tmp and os.path.exists(tmp):
                                    os.remove(tmp)
                                return True

                if is_tmp and os.path.exists(tmp):
                    os.remove(tmp)

            elif ext in ['.txt', '.csv']:
                # 读取文本文件的前几行
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    for i, line in enumerate(f):
                        if i >= 10:  # 只检查前10行
                            break
                        for pattern in target_patterns:
                            if pattern in line:
                                return True
        except Exception as e:
            # 静默失败，不影响主流程
            pass
        return False

    # ==========================================================================
    # 保存传给 AI 的真实文本到本地 TXT
    # ==========================================================================
    def _save_prompt_log(self, city, year, stage, system_msg, user_msg):
        """将每次上传给AI的文本保存为TXT，方便人工检查"""
        if not PROMPT_LOG_DIR:
            return

        # 构造文件名：城市_年份_阶段_时间戳.txt
        timestamp = datetime.now().strftime("%H%M%S")
        safe_city = re.sub(r'[\\/:*?"<>|]', '', city)
        file_name = f"{safe_city}_{year}_{stage}_{timestamp}.txt"
        file_path = os.path.join(PROMPT_LOG_DIR, file_name)

        # 拼接完整内容
        full_log_content = f"==========【System Prompt (系统指令)】==========\n{system_msg}\n\n"
        full_log_content += f"==========【User Prompt (实际发给AI的数据)】==========\n{user_msg}\n"

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(full_log_content)
        except Exception as e:
            print(f"   保存文本日志失败：{e}")

    # ==========================================================================
    # 检查文件是否包含排除关键词（如全市）
    # ==========================================================================
    def _is_file_excluded(self, file_path):
        """修改：仅检查文件是否完全是全市相关，不再直接排除文件，而是在读取时忽略内部的全市工作表"""
        file_name = os.path.basename(file_path).lower()
        # 仅排除文件名完全以"全市"开头且无市本级/市级关键词的文件
        if file_name.startswith("全市") and not any(kw.lower() in file_name for kw in CITY_LEVEL_KEYWORDS):
            return True
        return False

    # ==========================================================================
    # 获取文件内容中的市本级/市级关键词
    # ==========================================================================
    def _get_content_city_level_kw(self, file_path):
        """获取文件内容中包含的是市本级还是市级关键词"""
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, sheet_name=None, nrows=3, dtype=str)
                for sheet_name, sheet_df in df.items():
                    for _, row in sheet_df.head(3).iterrows():
                        row_text = ' '.join([str(val) for val in row.values if pd.notna(val)])
                        if "市本级" in row_text:
                            return "市本级"
                        elif "市级" in row_text:
                            return "市级"
            elif ext in ['.docx', '.doc']:
                tmp = file_path
                is_tmp = False
                if ext == '.doc' and Dispatch:
                    tmp = self._doc2docx(file_path)
                    is_tmp = True
                if tmp:
                    doc = docx.Document(tmp)
                    for i, para in enumerate(doc.paragraphs[:10]):
                        if "市本级" in para.text:
                            if is_tmp and os.path.exists(tmp):
                                os.remove(tmp)
                            return "市本级"
                        elif "市级" in para.text:
                            if is_tmp and os.path.exists(tmp):
                                os.remove(tmp)
                            return "市级"
                    if is_tmp and os.path.exists(tmp):
                        os.remove(tmp)
            elif ext in ['.txt', '.csv']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    for i, line in enumerate(f):
                        if i >= 10:
                            break
                        if "市本级" in line:
                            return "市本级"
                        elif "市级" in line:
                            return "市级"
        except:
            pass
        return ""

    # ==========================================================================
    # 新增：清洗内容中的全市相关文本
    # ==========================================================================
    def _clean_content_from_quanShi(self, content):
        """移除内容中所有包含「全市」的行/文本片段，确保提取阶段完全忽略"""
        if not content:
            return ""

        # 按行分割内容，逐行过滤
        lines = content.split('\n')
        cleaned_lines = []

        for line in lines:
            # 跳过包含"全市"的行
            if "全市" in line:
                continue
            # 移除行内的"全市"相关文本（如"全市金额"→"金额"）
            cleaned_line = re.sub(r'全市[\u4e00-\u9fa5]*', '', line)
            cleaned_lines.append(cleaned_line)

        # 重新拼接内容，清理多余空行
        cleaned_content = '\n'.join(cleaned_lines)
        cleaned_content = re.sub(r'\n{3,}', '\n\n', cleaned_content)

        return cleaned_content

    # ==========================================================================
    # 工作表筛选排序函数（增强版：考虑城市名）
    # ==========================================================================
    def _filter_and_sort_sheets(self, excel_path, city_name=""):
        """
        增强版工作表筛选排序：
        1. 忽略含"全市"等排除关键词的工作表
        2. 优先排序含城市名+（市本级/本级）的工作表
        3. 其次排序含城市名+（市级）的工作表
        4. 只包含城市名没有级别关键词的工作表忽略
        5. 返回排序后的工作表名称列表（保留所有有效工作表）
        """
        try:
            # 获取所有工作表名称
            excel_file = pd.ExcelFile(excel_path)
            all_sheets = excel_file.sheet_names

            city_core = city_name.replace("市", "") if city_name else ""

            # 分类工作表
            shi_benji_sheets = []  # 市本级工作表
            shi_ji_sheets = []  # 市级工作表
            exclude_sheets = []  # 排除的工作表
            other_sheets = []  # 其他工作表（无关键词）

            for sheet_name in all_sheets:
                sheet_name_lower = sheet_name.lower()

                # 第一步：检查是否为排除工作表
                if any(kw.lower() in sheet_name_lower for kw in SHEET_PRIORITY_KEYWORDS["排除"]):
                    exclude_sheets.append(sheet_name)
                    print(f"     忽略全市相关工作表: {sheet_name}")
                    continue

                # 第二步：检查是否为有效的市本级工作表（包含城市名+市本级/本级）
                is_valid_benji = False
                for pattern in [f"{city_core}市本级", f"{city_name}市本级", f"{city_core}本级", f"{city_name}本级"]:
                    if pattern.lower() in sheet_name_lower:
                        is_valid_benji = True
                        break

                if is_valid_benji:
                    shi_benji_sheets.append(sheet_name)
                    print(f"     优先读取市本级工作表: {sheet_name}")
                    continue

                # 第三步：检查是否为有效的市级工作表（包含城市名+市级）
                is_valid_ji = False
                for pattern in [f"{city_core}市级", f"{city_name}市级"]:
                    if pattern.lower() in sheet_name_lower:
                        is_valid_ji = True
                        break

                if is_valid_ji:
                    shi_ji_sheets.append(sheet_name)
                    print(f"     次优先读取市级工作表: {sheet_name}")
                    continue

                # 第四步：检查是否只包含城市名但没有级别关键词
                if city_core and city_core in sheet_name_lower:
                    has_level = any(kw.lower() in sheet_name_lower for kw in CITY_LEVEL_KEYWORDS)
                    if not has_level:
                        print(f"     忽略仅包含城市名的工作表: {sheet_name}")
                        continue

                # 其他工作表（保留，作为补充）
                other_sheets.append(sheet_name)
                print(f"     读取其他工作表: {sheet_name}")

            # 最终排序：市本级 > 市级 > 其他（无关键词）
            sorted_sheets = shi_benji_sheets + shi_ji_sheets + other_sheets

            # 打印筛选结果
            if exclude_sheets:
                print(f"     本次忽略工作表: {', '.join(exclude_sheets)}")
            if not sorted_sheets:
                print(f"     该Excel文件无有效工作表")

            return sorted_sheets

        except Exception as e:
            print(f"   筛选工作表出错 {excel_path}: {e}")
            # 出错时返回所有工作表（降级处理）
            try:
                return pd.ExcelFile(excel_path).sheet_names
            except:
                return []

    # ==========================================================================
    # 读取文件内容（优先处理指定工作表，保留多工作表文件完整）
    # ==========================================================================
    def read_file_content(self, path, city_name=""):
        ext = os.path.splitext(path)[1].lower()
        content = ""
        try:
            if ext in ['.xlsx', '.xls']:
                # 第一步：筛选并排序工作表（保留所有有效工作表）
                sorted_sheets = self._filter_and_sort_sheets(path, city_name)
                if not sorted_sheets:
                    return ""

                parts = []
                file_name = os.path.basename(path)
                parts.append(f"\n===== 文件：{file_name} =====")

                # 第二步：按优先级读取所有有效工作表（不再限制数量，保留完整内容）
                for sn in sorted_sheets:
                    # 读取单个工作表
                    df = pd.read_excel(path, sheet_name=sn, dtype=str)

                    # 清理空值和无关字符
                    df = df.fillna("").replace("\n", " ", regex=True)

                    cols = list(df.columns)
                    note_cols = []
                    # 标记列类型
                    for c in cols:
                        s = str(c).lower()
                        if any(kw.lower() in s for kw in INVALID_BUD_COLS):
                            note_cols.append(f"【预算列·禁止提取】{c}")
                            df[c] = "【预算数据·已屏蔽】"
                        # 新增：屏蔽含"全市"的列
                        elif "全市" in str(c):
                            note_cols.append(f"【全市列·禁止提取】{c}")
                            df[c] = "【全市数据·已屏蔽】"
                        elif any(kw.lower() in s for kw in VALID_DEC_COLS):
                            note_cols.append(f"【决算列·可提取】{c}")
                        else:
                            note_cols.append(f"【普通列】{c}")

                    # 转换为清晰的文本格式
                    txt = df.to_csv(sep='|', na_rep='无数据', index=False)
                    parts.append(f"--- 工作表: {sn} ---")
                    parts.append(f"列说明：{'; '.join(note_cols)}")
                    parts.append(txt)

                content = "\n".join(parts)

            elif ext in ['.docx', '.doc']:
                tmp = path
                is_tmp = False
                if ext == '.doc' and Dispatch:
                    tmp = self._doc2docx(path)
                    is_tmp = True
                if not tmp:
                    return ""
                doc = docx.Document(tmp)
                file_name = os.path.basename(path)

                lines = []
                lines.append(f"\n===== 文件：{file_name} =====")

                for p in doc.paragraphs:
                    t = p.text.strip()
                    # 过滤含"全市"的段落
                    if "全市" in t:
                        continue
                    if t and len(t) > 2:  # 过滤太短的无效文本
                        lines.append(t)
                for tb in doc.tables:
                    lines.append("\n[表格开始]")
                    headers = [c.text.strip() for c in tb.rows[0].cells] if tb.rows else []
                    h_note = []
                    for h in headers:
                        s = h.lower()
                        if any(kw.lower() in s for kw in INVALID_BUD_COLS):
                            h_note.append(f"【预算列·禁止提取】{h}")
                        # 新增：标记含"全市"的列
                        elif "全市" in h:
                            h_note.append(f"【全市列·禁止提取】{h}")
                        elif any(kw.lower() in s for kw in VALID_DEC_COLS):
                            h_note.append(f"【决算列·可提取】{h}")
                        else:
                            h_note.append(f"【普通列】{h}")
                    lines.append("表头说明：" + "; ".join(h_note))
                    # 只保留有数值的行，过滤含"全市"的行
                    for r in tb.rows:
                        row_v = [c.text.strip().replace('\n', ' ') for c in r.cells]
                        # 跳过包含"全市"的行
                        if any("全市" in cell for cell in row_v):
                            continue
                        if any(row_v) and len(row_v) == len(headers):
                            lines.append(" | ".join(row_v))
                    lines.append("[表格结束]\n")
                content = "\n".join(lines)
                if is_tmp and os.path.exists(tmp):
                    os.remove(tmp)

            elif ext in ['.txt', '.csv']:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    file_name = os.path.basename(path)
                    content = f"\n===== 文件：{file_name} =====\n" + content
                    # 清理多余空行
                    content = re.sub(r'\n{3,}', '\n\n', content)
        except Exception as e:
            print(f"   读取文件出错 {path}: {e}")
            return ""

        # 新增：清洗内容中的全市相关文本
        content = self._clean_content_from_quanShi(content)

        # 最终内容长度限制，避免超出API上下文（适当放宽到80000字符）
        if len(content) > 80000:
            content = content[:80000] + "\n[该文件内容过长，已截断核心部分（优先保留市本级/市级工作表）]"

        return content

    # ==========================================================================
    # 🔴 最终版文件过滤：完全按你的要求实现
    # 提取阶段准入 + 优先级排序
    # ==========================================================================
    def filter_and_sort_files(self, file_list, city, year):
        """
        你的规则 100% 实现：
        1. 必须提取：
           - 含 城市名+（市本级/本级/市级） 的所有文件
           - 所有多表同工文件
           - 所有多工作表文件
        2. 优先级：
           1) 含关键词 的 多表同工/多工作表
           2) 不含关键词 的 多表同工/多工作表
           3) 含关键词 的普通单表
        3. 排除：
           既不是多表同工/多工作表，也不含关键词 → 彻底排除
        """
        # 基础过滤
        valid = [f for f in file_list if not self._is_file_excluded(f)]
        valid = [f for f in valid if os.path.splitext(f)[1].lower() in SUPPORTED_EXTS]
        if self.target_keywords:
            valid = [f for f in valid if any(kw in os.path.basename(f) for kw in self.target_keywords)]

        # --- 分组：严格按你的重要性排序 ---
        group1 = []  # 最高：含关键词 + 多表同工 / 多工作表
        group2 = []  # 次之：不含关键词 + 多表同工 / 多工作表
        group3 = []  # 再次：含关键词 + 普通单表

        for f in valid:
            is_special = self._is_special_important_file(f)
            has_key = self._has_city_keyword(f, city)

            if is_special and has_key:
                group1.append(f)
                print(f"     优先【含关键词+多表/多工作表】: {os.path.basename(f)}")
            elif is_special:
                group2.append(f)
                print(f"     纳入【多表/多工作表】: {os.path.basename(f)}")
            elif has_key:
                group3.append(f)
                print(f"     纳入【含关键词单表】: {os.path.basename(f)}")
            else:
                print(f"     已排除【普通非市级文件】: {os.path.basename(f)}")

        # 最终排序
        final_files = group1 + group2 + group3
        final_files = list(dict.fromkeys(final_files))  # 去重

        if not final_files:
            print(f"     {city}{year}年无符合提取条件的文件")
        return final_files

    def _doc2docx(self, path):
        try:
            word = Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0  # 禁用弹窗
            doc = word.Documents.Open(path)
            out = path + "x"
            doc.SaveAs2(out, FileFormat=12)
            doc.Close()
            word.Quit()
            return out
        except Exception as e:
            print(f"   DOC转DOCX失败：{e}")
            return None

    # ==========================================================================
    # 增强的数值验证函数
    # ==========================================================================
    def _validate_numeric_value(self, value):
        """严格验证数值的有效性，避免混淆"""
        if not value or value == "未找到":
            return "未找到"

        # 转换为字符串处理
        val_str = str(value).strip()

        # 移除千分位逗号、单位等干扰字符
        val_str = re.sub(r',', '', val_str)  # 移除逗号
        val_str = re.sub(r'[万元亿]+$', '', val_str)  # 移除末尾的单位

        # 只保留数字、小数点和负号（处理负数）
        val_str = re.sub(r'[^\d.-]', '', val_str)

        # 验证是否为有效数值
        try:
            # 尝试转换为数值类型
            num = float(val_str)
            # 数值不能为负（决算数据应为非负）
            if num < 0:
                return "未找到"
            # 返回标准化的数值字符串
            return str(num)
        except:
            return "未找到"

    # ==========================================================================
    # 单表优先提取核心逻辑（增强版）
    # ==========================================================================
    def analyze_single_year(self, city, year, files):
        """核心规则：优先从第一张核心表提取所有指标，不足时再补充（简化输出）"""
        # 第一步：先尝试从第一张核心表提取所有指标
        primary_result = None
        if files:
            # 对files按优先级排序：一般公共预算本级支出表优先
            files = self._sort_files_by_priority(files)
            primary_file = files[0]  # 取优先级最高的核心表

            # 读取核心表内容（包含所有有效工作表）
            primary_content = self.read_file_content(primary_file, city)
            if primary_content:
                primary_result = self._extract_from_single_file(city, year, primary_content, primary_file)

                # 检查是否所有指标都已找到
                found_all = all([primary_result[p] != "未找到" for p in PARAMETERS])
                if found_all:
                    primary_result["数据来源"] = os.path.basename(primary_file)
                    primary_result["是否单表提取"] = "是"
                    return primary_result

        # 第二步：核心表未找全，整合所有文件补充
        full_context = ""
        file_names = []
        for f in files[:8]:  # 放宽到8个文件，充分利用多工作表文件的信息
            content = self.read_file_content(f, city)
            if content:
                full_context += content
                file_names.append(os.path.basename(f))

        if not full_context.strip():
            res = {"年份": year, "数据来源": "无有效文件", "是否单表提取": "否"}
            for p in PARAMETERS:
                res[p] = "未找到"
            return res

        # 调用AI补充提取
        final_result = self._extract_from_multiple_files(city, year, full_context, primary_result, file_names)
        final_result["数据来源"] = "; ".join(file_names[:3]) + ("..." if len(file_names) > 3 else "")
        final_result["是否单表提取"] = "否"

        return final_result

    def _sort_files_by_priority(self, files):
        """
        对文件按优先级排序：
        1. 一般公共预算本级支出表（最高优先级）
        2. 其他包含"市本级"的文件
        3. 其他包含"市级"的文件
        4. 其他文件
        """

        def get_priority(file_path):
            fname = os.path.basename(file_path).lower()
            # 最高优先级：一般公共预算本级支出表
            if "一般公共预算本级支出表" in fname:
                return 0
            # 第二优先级：市本级相关
            if "市本级" in fname:
                return 1
            # 第三优先级：市级相关
            if "市级" in fname:
                return 2
            # 其他
            return 3

        return sorted(files, key=get_priority)

    def _extract_from_single_file(self, city, year, content, file_name):
        """从单文件提取所有指标（增强版，精准区分相似指标）"""
        system_msg = f"""你是专业的财政数据提取专家，严格遵守以下规则：

【基础信息】
- 当前城市：{city}
- 当前年份：{year}
- 数据文件：{file_name}

【核心规则 - 必须100%遵守】
1. 严格区分以下指标，绝对不能混淆：
   - "财政事务"：指该大类的总决算金额
   - "财政事务——信息化建设"：仅指财政事务中信息化建设专项的决算金额
   - "税收事务"：指该大类的总决算金额
   - "税收事务——信息化建设"：仅指税收事务中信息化建设专项的决算金额
   - "审计事务"：指该大类的总决算金额
   - "审计事务——信息化建设"：仅指审计事务中信息化建设专项的决算金额
2. 只提取{year}年的市本级/市级一般公共预算支出决算数据，严禁跨年份
3. 只从【决算列·可提取】列中提取数值，【预算列·禁止提取】和【全市列·禁止提取】列的数据绝对不提取
4. 完全忽略任何包含"全市"字样的行、列、数值，绝对不能使用这些数据
5. 只提取有明确、具体数值的内容，没有明确数值的填"未找到"
6. 严禁无中生有、猜测数值，找不到的指标必须填"未找到"
7. 提取的数值仅保留数字和小数点，移除单位（如万元、亿）和千分位逗号
8. 信息化建设专项金额通常小于对应大类总金额，如果发现相反情况，判定为"未找到"
9. 优先从标注为"{city}市本级"或"{city}本级"的工作表提取数据，其次是"{city}市级"工作表

【提取指标清单】
{chr(10).join(PARAMETERS)}

【输出要求】
- 仅返回标准JSON格式，无任何多余文字
- JSON中必须包含所有指标，值为找到的具体数值（仅数字）或"未找到"
- 必须包含"年份"字段，值为{year}
"""
        messages = [
            {"role": "system", "content": system_msg},
            {"role": "user",
             "content": f"请从以下{year}年数据中精准提取指定指标，优先使用{city}市本级工作表数据，严格区分相似指标，完全忽略所有含'全市'的内容：\n{content}"}
        ]

        # 保存AI交互文本
        self._save_prompt_log(city, year, "1_核心表提取", system_msg, messages[1]["content"])

        try:
            resp = self.client.chat.completions.create(
                model=QWEN_MODEL,
                messages=messages,
                temperature=0.0,  # 零温度确保结果稳定
                timeout=150,
                max_tokens=2000
            )
            txt = resp.choices[0].message.content.strip()
            result = self._parse_safe_json(txt, year)

            # 增强的数据验证
            validated_result = self._validate_extracted_data(result)

            # 额外的逻辑验证：信息化专项不能大于大类总额
            validated_result = self._validate_hierarchy_logic(validated_result)

            return validated_result
        except Exception as e:
            print(f"   单文件提取失败：{str(e)[:50]}")
            res = {"年份": year}
            for p in PARAMETERS:
                res[p] = "未找到"
            return res

    def _extract_from_multiple_files(self, city, year, content, primary_result, file_names):
        """整合多文件补充提取缺失指标（增强版，保护已正确提取的值）"""
        missing_hint = ""
        found_params = []
        missing_params = []

        if primary_result:
            found_params = [p for p in PARAMETERS if primary_result[p] != "未找到"]
            missing_params = [p for p in PARAMETERS if primary_result[p] == "未找到"]

            found_str = ", ".join([f"{p}: {primary_result[p]}" for p in found_params])
            missing_str = ", ".join(missing_params)

            missing_hint = f"""
【已找到的指标】（以下值已验证正确，严禁修改，仅补充缺失指标）：
{found_str}

【需要补充的指标】（仅查找以下指标，不要修改已找到的指标）：
{missing_str}
"""
        system_msg = f"""你是专业的财政数据提取专家，严格遵守以下规则：

【基础信息】
- 当前城市：{city}
- 当前年份：{year}
- 数据源：{', '.join(file_names)}

【核心规则 - 必须100%遵守】
1. 严格区分以下指标，绝对不能混淆：
   - "财政事务"：指该大类的总决算金额
   - "财政事务——信息化建设"：仅指财政事务中信息化建设专项的决算金额
   - "税收事务"：指该大类的总决算金额
   - "税收事务——信息化建设"：仅指税收事务中信息化建设专项的决算金额
   - "审计事务"：指该大类的总决算金额
   - "审计事务——信息化建设"：仅指审计事务中信息化建设专项的决算金额
2. 补充的指标必须是{year}年的市本级/市级决算数据，严禁跨年份、跨城市
3. 只从【决算列·可提取】列中提取数值，预算数据和全市数据绝对不提取
4. 完全忽略任何包含"全市"字样的行、列、数值，绝对不能使用这些数据
5. 只提取有明确数值的内容，找不到的仍填"未找到"
6. 严禁无中生有、猜测数值，确保所有数值真实存在
7. 提取的数值仅保留数字和小数点，移除单位（如万元、亿）和千分位逗号
8. 信息化建设专项金额通常小于对应大类总金额，如果发现相反情况，判定为"未找到"
9. 优先从标注为"{city}市本级"或"{city}本级"的工作表提取数据，其次是"{city}市级"工作表
10. 已找到的指标值必须完全保留，仅补充缺失指标的值

{missing_hint}

【提取指标清单】
{chr(10).join(PARAMETERS)}

【输出要求】
- 仅返回标准JSON格式，包含所有指标
- 已找到的指标保留原值，缺失的指标补充后的值或"未找到"
- 必须包含"年份"字段，值为{year}
"""
        messages = [
            {"role": "system", "content": system_msg},
            {"role": "user",
             "content": f"请补充提取以下{year}年数据中的缺失指标，优先使用{city}市本级工作表数据，严格区分相似指标，完全忽略所有含'全市'的内容：\n{content}"}
        ]

        # 保存AI交互文本
        self._save_prompt_log(city, year, "2_补充文件提取", system_msg, messages[1]["content"])

        try:
            resp = self.client.chat.completions.create(
                model=QWEN_MODEL,
                messages=messages,
                temperature=0.0,
                timeout=150,
                max_tokens=2000
            )
            txt = resp.choices[0].message.content.strip()
            result = self._parse_safe_json(txt, year)

            # 合并已有结果和补充结果（保护已正确提取的值）
            if primary_result:
                for p in PARAMETERS:
                    if primary_result[p] != "未找到":
                        result[p] = primary_result[p]

            result = self._validate_extracted_data(result)
            result = self._validate_hierarchy_logic(result)

            return result
        except Exception as e:
            print(f"   多文件补充提取失败：{str(e)[:50]}")
            if primary_result:
                return primary_result
            else:
                res = {"年份": year}
                for p in PARAMETERS:
                    res[p] = "未找到"
                return res

    def _validate_hierarchy_logic(self, data):
        """验证层级逻辑：信息化专项金额不能大于对应大类金额"""
        hierarchy_pairs = [
            ("财政事务", "财政事务——信息化建设"),
            ("税收事务", "税收事务——信息化建设"),
            ("审计事务", "审计事务——信息化建设")
        ]

        for main_param, sub_param in hierarchy_pairs:
            main_val = data.get(main_param, "未找到")
            sub_val = data.get(sub_param, "未找到")

            if main_val != "未找到" and sub_val != "未找到":
                try:
                    main_num = float(main_val)
                    sub_num = float(sub_val)
                    if sub_num > main_num:
                        data[sub_param] = "未找到"
                except:
                    pass
        return data

    def _validate_extracted_data(self, data):
        """增强版数据验证，过滤不合理值并标准化"""
        data["年份"] = str(data.get("年份", "")).strip()

        for param in PARAMETERS:
            if param not in data:
                data[param] = "未找到"
                continue
            data[param] = self._validate_numeric_value(data[param])
        return data

    def _parse_safe_json(self, txt, year):
        """安全解析JSON，增强容错"""
        try:
            txt = re.sub(r'```json|```', '', txt).strip()
            match = re.search(r'\{[\s\S]*\}', txt)
            if not match:
                raise Exception("未找到JSON结构")

            json_str = match.group()
            json_str = re.sub(r"'", '"', json_str)
            json_str = re.sub(r',\s*}', '}', json_str)

            d = json.loads(json_str)
            d["年份"] = year
            for p in PARAMETERS:
                if p not in d:
                    d[p] = "未找到"
            return d
        except Exception as e:
            print(f"   JSON解析失败：{str(e)[:30]}")
            d = {"年份": year}
            for p in PARAMETERS:
                d[p] = "未找到"
            return d

    # ==========================================================================
    # 城市主逻辑（简化输出+精简Excel列）
    # ==========================================================================
    def process_single_city(self, city):
        """处理单个城市，简化终端输出+精简Excel列"""
        if city in self.breakpoints["completed"]:
            print(f"   {city} 已处理过，跳过")
            return
        if self._is_city_skipped(city):
            print(f"   {city} 标记为跳过，跳过")
            return

        path = self._get_city_path(city)
        if not path or not os.path.exists(path):
            print(f"   {city} 路径不存在：{path}")
            self._save_breakpoint(city)
            return

        year_map = {}
        for d in os.listdir(path):
            dp = os.path.join(path, d)
            if os.path.isdir(dp) and re.match(r'^20\d{2}$', d):
                files = []
                for r, _, fs in os.walk(dp):
                    for f in fs:
                        files.append(os.path.join(r, f))
                year_map[d] = files

        if not year_map:
            print(f"   {city} 无年份文件夹")
            self._save_breakpoint(city)
            return

        results = []
        for year in sorted(year_map.keys()):
            print(f"   处理 {city} {year} 年")
            # 传入city和year参数
            files = self.filter_and_sort_files(year_map[year], city, year)
            res = self.analyze_single_year(city, year, files)
            res["城市"] = city
            res["省名"] = self.city_config_df[self.city_config_df["市名"] == city].iloc[0]["省名"]

            res.pop("是否单表提取", None)
            res.pop("数据来源", None)

            results.append(res)

        if results:
            df = pd.DataFrame(results)
            if os.path.exists(OUTPUT_EXCEL_PATH):
                old_df = pd.read_excel(OUTPUT_EXCEL_PATH, dtype=str).fillna("")
                old_df = old_df[~((old_df["城市"] == city) & (old_df["年份"].astype(str).isin(year_map.keys())))]
                df = pd.concat([old_df, df], ignore_index=True)

            cols_order = ["省名", "城市", "年份"] + PARAMETERS
            df = df.reindex(columns=cols_order, fill_value="")

            df.to_excel(OUTPUT_EXCEL_PATH, index=False)
            print(f"   已保存 {city} 数据")

        self._save_breakpoint(city)

    # ==========================================================================
    # 入口
    # ==========================================================================
    def select_data_scope(self):
        print("\n1:决算 2:决算报告 3:预算 4:预算报告 5:全部")
        sel = input("请输入序号: ").replace('，', ' ').replace(',', ' ').split()
        mp = {"1": "决算", "2": "决算报告", "3": "预算", "4": "预算报告"}
        self.target_keywords = [mp[s] for s in sel if s in mp]
        if "5" in sel or not self.target_keywords:
            self.target_keywords = []
        print(f"已选：{'全部' if not self.target_keywords else ' + '.join(self.target_keywords)}")

    def run(self):
        self.select_data_scope()
        if self.city_config_df.empty or not self.city_order_list:
            print("配置文件为空或无城市列表")
            return
        print(f"\n开始处理 {len(self.city_order_list)} 个城市")
        print(f"实际上传给 AI 的数据将被备份至：{PROMPT_LOG_DIR}")
        print(f"核心表将被备份至：{PRIMARY_FILES_SAVE_DIR}")
        print(f"已启用过滤规则：读取文件时忽略全市工作表，提取指标时完全忽略所有含'全市'的内容")
        print(f"已按你要求排序：含关键词多表 > 普通多表 > 含关键词单表")
        print(f"单表优先规则：一般公共预算本级支出表 > 市本级相关表 > 市级相关表")
        print(f"多表同工规则：优先选择包含城市名+市本级/市级/本级的表，只包含城市名的表忽略")
        for i, city in enumerate(self.city_order_list, 1):
            print(f"\n===== {i}/{len(self.city_order_list)} {city} =====")
            self.process_single_city(city)
        print("\n全部完成")


# ==============================================================================
# 运行入口
# ==============================================================================
if __name__ == "__main__":
    analyzer = RobustBudgetAnalyzer()
    analyzer.run()
    input("\n按回车退出程序...")